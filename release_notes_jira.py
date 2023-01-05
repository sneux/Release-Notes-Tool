import requests
import json
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm, RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement, qn

def get_lists(version: str) -> str:
  with open("AUTH.json", 'r') as file:
    contents = json.load(file)
    user_auth = contents["Authorization"]

  version = version
  url = f"http://dataremote.atlassian.net/rest/api/2/search?jql=project in (CXX,VOIP,DEV) AND fixVersion in (\"{version}\" ) ORDER BY project, updated DESC"
  payload = ""
  headers = {"Authorization": f"Basic {user_auth}"}
  response = requests.request("GET", url, headers=headers, data=payload)
  data = json.loads(response.text)          # all of the data inside the JSON file
  issues = data['issues']                   # a list of dictionaries containing the info we need
  bug_list = []
  other_list = []

  # we have to iterate over the list of dictionaries and manually find the info we need
  index = 0
  for item in issues:
    key = (issues[index]['key'])
    type = (issues[index]['fields']['issuetype']['name'])
    
    # we have to first see if the custom fields exist for the project and if not
    # we pass the errror it raises to keep iterating 
    if type == 'Bug':
      try:
        release_n = issues[index]['fields']['customfield_10691']
      except (KeyError):
        try:
          release_n = issues[index]['fields']['customfield_10615']
        except (KeyError):
          pass 
      bug_list.append((key,release_n))
    
    else:
      try:
        release_n = issues[index]['fields']['customfield_10691']
      except (KeyError):
        try:
          release_n = issues[index]['fields']['customfield_10615']
        except (KeyError):
          release_n = "N/A"  
      other_list.append((key,release_n))
    index += 1
  return other_list, bug_list

# we need to get rid of all the "N/A" entries
def clean_list(a_list):
  test_string = "N/A"
  new_list = [item for item in a_list if test_string not in item]
  return new_list

# sets the background of a cell to a specific hex code 
# the make_table method (directly below) calls this method
def fill_cell_bg(table, color, index):
  row = table.rows[0].cells
  cell = row[index]
  bg = cell._tc.get_or_add_tcPr()
  hAlign = OxmlElement("w:shd")
  hAlign.set(qn("w:fill"), color)
  bg.append(hAlign)

def make_table(header, a_list):
  # we have to create a table just for the header 
  table = doc.add_table(rows = 1, cols = 1, style = "Table Grid")
  table_header = table.rows[0].cells
  run = table_header[0].paragraphs[0].add_run(header)
  run.bold = True
  table_header[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  cell_bg = fill_cell_bg(table, "#C00000", 0)

  # now we create a table for the ID and Description which will automatically 
  # attach to our header table
  table2 = doc.add_table(rows = 1, cols = 2, style = "Table Grid")
  table2.allow_autofit = False
  table2.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

  # title the ID and Description cells 
  row = table2.rows[0].cells
  row[0].text = "ID#"
  cell_bg2 = fill_cell_bg(table2, "#262626", 0)
  row[1].text = "Description"
  cell_bg3 = fill_cell_bg(table2, "#262626", 1)

  # populate the table with our release notes
  for id in a_list:
    row = table2.add_row().cells
    row[0].text = id[0]
    row[1].text = id[1]
  for cell in table2.columns[1].cells:
      cell.width = Inches(10)

def add_header(header):
  section_header = doc.add_heading('')
  section_header.add_run(header).font.color.rgb = RGBColor(0,0,0)
  rFonts = section_header.style.element.rPr.rFonts
  rFonts.set(qn("w:asciiTheme"), "Times New Roman")

def add_title(title):
  p = doc.add_paragraph()
  header = doc.add_paragraph('')
  header_title = header.add_run(title)
  header_title.font.color.rgb = RGBColor(0,0,0)
  header.alignment = 1
  header_title.font.size = Pt(20)
  header_title.bold = True

def make_doc(other_list, bug_list):
  #first let's create and format the first page 
  global doc
  doc = Document()
  style = doc.styles['Normal']
  font = style.font
  font.name = "Times New Roman"
  font.size = Pt(12)

  # create white space so our cover page starts in the middle
  for x in range(5):
    p = doc.add_paragraph()
  
  # add our data remote logo
  doc.add_picture('dataremotelogo.png')
  last_paragraph = doc.paragraphs[-1]
  last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  # now lets add the titles to our cover page
  header = add_title("VAB-1")
  header2 = add_title("<FIRMWARE Version>")
  header3 = add_title("Release Notes")
  doc.add_page_break()

  # adding the header to each page
  header_img = doc.sections[0].header
  paragraph = header_img.paragraphs[0]
  logo_header = paragraph.add_run()
  logo_header.add_picture("dataremotelogo.png", width=Inches(3))

  # make it so the header image doesn't show on the first page
  for section in doc.sections:
    section.different_first_page_header_footer = True
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

  # now we add the software new features section
  new_features_header = add_header("Software New Features")
  nf_desc = doc.add_paragraph('This section identifies new software features on this release.')
  # creating the table
  table = make_table("New Features", other_list)
  doc.add_page_break()

  # and the page with "known issues" table
  resolved_issues_header = add_header("Software Resolved Issues")
  ri_desc = doc.add_paragraph('This section identifies issues that have been resolved since the last release of the software.')
  table2 = make_table("Resolved Issues", bug_list)

  doc.save('Release Notes.docx')

def main():
  other_list, bug_list = get_lists("ATT_R3.1")
  bug_list = clean_list(bug_list)
  other_list = clean_list(other_list)
  make_doc(other_list, bug_list)

if __name__ == "__main__":
    main()