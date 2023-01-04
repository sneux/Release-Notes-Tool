import pandas as pd
from pandas._libs.parsers import STR_NA_VALUES
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm, RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement, qn

# method that returns a list of multiples of the same column name 
# this is a helper function so we can combine multiple release_columns into one 
# THIS METHOD IS CALLED BY THE COMBINE_RELEASE_COLUMNS METHOD DIRECTLY BELOW
def get_release_columns(file_name):
    df = pd.read_csv(file_name)                 # create a dataframe from a csv file
    all_release_columns = list(df.columns.values)      # get a list of data frame column names

    test_string = "Release Notes"               # substring to search in list
    release_columns = []                        # new list to store indices containing our substring
    i = 0                                       
    list_length = len(all_release_columns)             

    # we need to iterate over the list and check if the substring "release notes" is in each item in the list
    # if there's a match, store the index in the release_columns list
    while i < list_length:
        if(all_release_columns[i].find(test_string) != -1):
            release_columns.append(all_release_columns[i])
        i += 1
    return(release_columns)

# a method to clean and combine columns with the same name and information 
# returns one Series containing the release notes 
# THIS METHOD IS CALLED BY COMBINE_DFS COLUMN DIRECTLY BELOW
def combine_release_columns(file_name):
    accepted_na_values = STR_NA_VALUES - {'N/A'} | {'_'}    # accept manual "N/A" entries as strings
                                                            # pandas interprets N/A cells as nan
    col_names = get_release_columns(file_name)              # list of release notes column names

    # create a new df that replaces all empty cells with string for easier removal later
    df = pd.read_csv(file_name, keep_default_na=False, na_values=accepted_na_values).fillna('empty space')
    # we pass our list of column names using list comprehension
    new_df = pd.concat([df.loc[:, f"{name}"] for name in col_names], ignore_index=False).filter(regex='^((?!empty space).)*$', axis=0)
    # filter out the empty space cells
    filtered_s = new_df[~new_df.str.contains('empty space')]
    
    return filtered_s

# a method combining our release notes series with our issue type and key dataframe
def combine_dfs(file_name):
    release_series = combine_release_columns(file_name)     # a single Series including our combined release notes
    df2 = pd.read_csv(file_name, usecols=[0,1])              # a datafrane containing only our issue type and issue key

    # merge our release notes into our data frame by index
    final_df = df2.merge(release_series.rename('Release Notes'), left_index=True, right_index=True)

    return final_df

# a method that returns two lists containing just the issue type and the corresponding release notes
def get_lists(df):
    # we use a list of tuples to store the key and the release notes 
    # we use two different lists because the information stored in each list is 
    # going to different places
    bug_list = list()
    fixed_list = list() 

    # first we need to get the issue type of the project to separate the bugs 
    for index,row in df.iterrows():
        issue_type = row[0]                 # store the name of our issue type, key and release notes for easier referencing
        issue_key = row[1]                  
        release_note = row[2]
        if issue_type == 'Bug':
            bug_list.append([issue_key, release_note])
        else:
            fixed_list.append([issue_key, release_note])

    return bug_list, fixed_list

def clean_list(a_list):
  test_string = "N/A"
  new_list = [item for item in a_list if test_string not in item]
  return new_list

# sets the background of a cell to a specific hex code 
# the make_table method (directly below) calls this method
def fill_cell_bg(table, color, index):
  row = table.rows[0].cells
  cell = row[index]
  bg = cell._tc.get_or_add_tcPr()
  hAlign = OxmlElement("w:shd")
  hAlign.set(qn("w:fill"), color)
  bg.append(hAlign)

def make_table(header, a_list):
  # we have to create a table just for the header 
  table = doc.add_table(rows = 1, cols = 1, style = "Table Grid")
  table_header = table.rows[0].cells
  run = table_header[0].paragraphs[0].add_run(header)
  run.bold = True
  table_header[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  cell_bg = fill_cell_bg(table, "#C00000", 0)

  # now we create a table for the ID and Description which will automatically 
  # attach to our header table
  table2 = doc.add_table(rows = 1, cols = 2, style = "Table Grid")
  table2.allow_autofit = False
  table2.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

  # title the ID and Description cells 
  row = table2.rows[0].cells
  row[0].text = "ID#"
  cell_bg2 = fill_cell_bg(table2, "#262626", 0)
  row[1].text = "Description"
  cell_bg3 = fill_cell_bg(table2, "#262626", 1)

  # populate the table with our release notes
  for id in a_list:
    row = table2.add_row().cells
    row[0].text = id[0]
    row[1].text = id[1]
  for cell in table2.columns[1].cells:
      cell.width = Inches(10)

def add_header(header):
  section_header = doc.add_heading('')
  section_header.add_run(header).font.color.rgb = RGBColor(0,0,0)
  rFonts = section_header.style.element.rPr.rFonts
  rFonts.set(qn("w:asciiTheme"), "Times New Roman")

def add_title(title):
  header = doc.add_heading('', 0)
  header.add_run(title).font.color.rgb = RGBColor(0,0,0)
  header.alignment = 1
  rFonts = header.style.element.rPr.rFonts
  rFonts.set(qn("w:asciiTheme"), "Times New Roman")

def make_doc(other_list, bug_list):
  #first let's create and format the first page 
  global doc
  doc = Document()

  # set the doc font
  # NOTE: this doesnt effect header fonts
  style = doc.styles['Normal']
  font = style.font
  font.name = "Times New Roman"
  font.size = Pt(12)

  # create white space so our cover page starts in the middle
  for x in range(5):
    p = doc.add_paragraph()
  
  # add our data remote logo
  doc.add_picture('dataremotelogo.png')
  last_paragraph = doc.paragraphs[-1]
  last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  # now lets add the titles to our cover page
  header = add_title("VAB-1")
  header2 = add_title("<FIRMWARE Version>")
  header3 = add_title("Release Notes")
  doc.add_page_break()

  # adding the header to each page
  header_img = doc.sections[0].header
  paragraph = header_img.paragraphs[0]
  logo_header = paragraph.add_run()
  logo_header.add_picture("dataremotelogo.png", width=Inches(3))

  # make it so the header image doesn't show on the first page
  for section in doc.sections:
    section.different_first_page_header_footer = True
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


  # now we add the software new features section
  new_features_header = add_header("Software New Features")
  nf_desc = doc.add_paragraph('This section identifies new software features on this release.')
  # creating the table
  table = make_table("New Features", other_list)
  doc.add_page_break()

  # and the page with "known issues" table
  resolved_issues_header = add_header("Software Resolved Issues")
  ri_desc = doc.add_paragraph('This section identifies issues that have been resolved since the last release of the software.')
  table2 = make_table("Resolved Issues", bug_list)

  doc.save('Release Notes CSV.docx')

def main():
    csv_file = '~\Desktop\VAB1.csv'
    final_df = combine_dfs(csv_file)
    bugs, fixedR = get_lists(final_df)
    bug_list = clean_list(bugs)
    other_list = clean_list(fixedR)
    make_doc(other_list, bug_list)

if __name__ == "__main__":
    main()
